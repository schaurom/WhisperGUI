# grafische Oberfläche zum Auswählen von Dateien

import tkinter as tk
from tkinter import filedialog, messagebox
import whisper
from docx import Document
from docx.shared import Pt, RGBColor
import time

def run_whisper(audio_file_path, output_file_path):
    # Startzeit für Programm-laufzeit
    startzeit = time.time()

    print("\n"+audio_file_path)
    # models: tiny base small medium large-v1 large-v2
    model = whisper.load_model("tiny", download_root='.\models')
    options = {"language": "de", "verbose": "true", "word_timestamps": "true", "append_punctuations": "."}
    result = model.transcribe(audio_file_path, **options)

    #Inhalt in Variable schreiben
    result_text = result["text"]
    print("\n"+result_text)
    result_details = result['segments']

    # Word
    dateiname = output_file_path

    # erstelle ein neues Word Objekt
    dokument = Document()
    absatz = dokument.add_paragraph()

    # Farben definieren
    GREEN = RGBColor(0, 128, 0)
    ORANGE = RGBColor(255, 165, 0)
    RED = RGBColor(255, 0, 0)

    # Klossar einfügen
    absatz.add_run('Dunkelgrün für hohe Genauigkeit - >80%').font.size = Pt(14)
    absatz.runs[-1].font.color.rgb = GREEN

    absatz.add_run('\nOrange für moderate Genauigkeit - 60% >= 80%').font.size = Pt(14)
    absatz.runs[-1].font.color.rgb = ORANGE

    absatz.add_run('\nRot für niedrige Genauigkeit - <60%').font.size = Pt(14)
    absatz.runs[-1].font.color.rgb = RED
    absatz.add_run('\n--------------------------------------------------------------------------------------\n\n\n').font.size = Pt(14)


    # alle Wörter durchiterieren
    for element in result_details:
        for listitem in element['words']:
            word_text = listitem['word']
            word_probability = listitem['probability']
            print(word_text)
            print(word_probability)

            # Wort in das Dokument einfügen, Schriftgröße spezifizieren und Farbe basierend auf der Genauigkeit
            absatz.add_run(word_text).font.size = Pt(14)
            if word_probability > 0.8:
                absatz.runs[-1].font.color.rgb = GREEN  # Dunkelgrün für hohe Genauigkeit
            elif word_probability > 0.6:
                absatz.runs[-1].font.color.rgb = ORANGE  # Orange für moderate Genauigkeit
            else:
                absatz.runs[-1].font.color.rgb = RED  # Rot für niedrige Genauigkeit

    
    # Speichern Sie das Dokument in einer Datei
    dokument.save(dateiname)
    print("\n"+dateiname)

    # Endzeit für Programm-laufzeit
    endzeit = time.time()
    # Die Laufzeit berechnen (in Sekunden)
    laufzeit = endzeit - startzeit
    print(f"\nLaufzeit des Programms: {laufzeit:.2f} Sekunden.\n")


def transcribe_audio():
    audio_file_path = filedialog.askopenfilename(filetypes=[("Audio files", "*.mp3")])
    if audio_file_path:
        output_file_path = audio_file_path.replace(".mp3", "_transcript.docx")
        try:
            # Transkription mit "whisper" durchführen
            run_whisper(audio_file_path, output_file_path)

            # Meldung anzeigen
            messagebox.showinfo("Erfolg", f"Transkription abgeschlossen. Ergebnis wurde in {output_file_path} gespeichert.")
        except Exception as e:
            print(str(e))
            messagebox.showerror("Fehler", f"Fehler bei der Transkription: {str(e)}")


# Hauptfenster erstellen
root = tk.Tk()
root.title("Whisper Transkription")
# Set the width and height of the main window
window_width = 400  # Set your desired width
window_height = 100  # Set your desired height
root.geometry(f"{window_width}x{window_height}")

# Button zum Auswählen der Audiodatei
select_file_button = tk.Button(root, text="Audiodatei auswählen und transkribieren", command=transcribe_audio)
select_file_button.pack(padx=20, pady=20)


# Tkinter Hauptloop starten
root.mainloop()
